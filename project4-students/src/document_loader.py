from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path
from typing import Iterable


@dataclass
class LoadedDocument:
    doc_id: str
    title: str
    source_path: str
    text: str


def _read_txt_like(path: Path) -> str:
    return path.read_text(encoding="utf-8", errors="ignore")


def _read_pdf(path: Path) -> str:
    try:
        from pypdf import PdfReader
    except ImportError as exc:
        raise ImportError("Install pypdf to parse PDF files: pip install pypdf") from exc

    reader = PdfReader(str(path))
    parts: list[str] = []
    for i, page in enumerate(reader.pages, start=1):
        page_text = page.extract_text() or ""
        if page_text.strip():
            parts.append(f"\n\n[Page {i}]\n{page_text}")
    return "\n".join(parts)


def _read_docx(path: Path) -> str:
    try:
        import docx
    except ImportError as exc:
        raise ImportError("Install python-docx to parse Word files: pip install python-docx") from exc

    doc = docx.Document(str(path))
    return "\n".join(p.text for p in doc.paragraphs if p.text.strip())


def load_documents(docs_dir: Path) -> list[LoadedDocument]:
    """Load .txt, .md, .pdf, and .docx documents from a directory."""
    supported = {".txt", ".md", ".pdf", ".docx"}
    docs: list[LoadedDocument] = []

    for path in sorted(docs_dir.iterdir()):
        if path.is_dir() or path.suffix.lower() not in supported:
            continue

        suffix = path.suffix.lower()
        if suffix in {".txt", ".md"}:
            text = _read_txt_like(path)
        elif suffix == ".pdf":
            text = _read_pdf(path)
        elif suffix == ".docx":
            text = _read_docx(path)
        else:
            continue

        cleaned = "\n".join(line.rstrip() for line in text.splitlines()).strip()
        if not cleaned:
            continue

        docs.append(
            LoadedDocument(
                doc_id=path.stem.lower().replace(" ", "_"),
                title=path.stem.replace("_", " ").title(),
                source_path=str(path.relative_to(docs_dir.parent.parent)),
                text=cleaned,
            )
        )

    return docs
